"""
Парсер документов для RAG-платформы
Поддерживает PDF, DOCX, XLSX, HTML, EML форматы
"""

import os
import hashlib
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass
import logging

# Импорты для различных форматов
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available, PDF parsing disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available, DOCX parsing disabled")

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    logging.warning("openpyxl not available, XLSX parsing disabled")

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    logging.warning("beautifulsoup4 not available, HTML parsing disabled")

try:
    import email
    EML_AVAILABLE = True
except ImportError:
    EML_AVAILABLE = False
    logging.warning("email module not available, EML parsing disabled")

@dataclass
class DocumentChunk:
    """Чанк документа"""
    text: str
    chunk_index: int
    page_no: Optional[int] = None
    table_html: Optional[str] = None
    bbox: Optional[Dict[str, Any]] = None
    metadata: Optional[Dict[str, Any]] = None

@dataclass
class DocumentContent:
    """Содержимое документа"""
    title: Optional[str]
    text: str
    tables: List[Dict[str, Any]]
    pages: List[Dict[str, Any]]
    sha256: str
    mime_type: str
    needs_ocr: bool
    metadata: Dict[str, Any]

class DocumentParser:
    """Основной парсер документов"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.xlsx': self._parse_xlsx,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.eml': self._parse_eml,
        }
    
    def parse(self, file_path: str) -> DocumentContent:
        """Парсинг документа по расширению файла"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Вычисляем SHA256
        sha256 = self._calculate_sha256(file_path)
        
        # Определяем MIME тип
        mime_type = self._get_mime_type(file_extension)
        
        # Парсим содержимое
        parse_func = self.supported_formats[file_extension]
        content = parse_func(file_path)
        
        # Определяем, нужен ли OCR
        needs_ocr = self._needs_ocr(file_extension, content)
        
        return DocumentContent(
            title=content.get('title'),
            text=content.get('text', ''),
            tables=content.get('tables', []),
            pages=content.get('pages', []),
            sha256=sha256,
            mime_type=mime_type,
            needs_ocr=needs_ocr,
            metadata=content.get('metadata', {})
        )
    
    def create_chunks(self, content: DocumentContent, chunk_size: int = 1000, overlap: int = 200) -> List[DocumentChunk]:
        """Создание чанков из текста документа"""
        chunks = []
        text = content.text
        
        if not text:
            return chunks
        
        # Простое разбиение по словам
        words = text.split()
        current_chunk = []
        current_length = 0
        chunk_index = 0
        
        for word in words:
            current_chunk.append(word)
            current_length += len(word) + 1  # +1 для пробела
            
            if current_length >= chunk_size:
                # Создаем чанк
                chunk_text = ' '.join(current_chunk)
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    chunk_index=chunk_index,
                    metadata={'chunk_size': current_length}
                ))
                
                # Начинаем новый чанк с overlap
                if overlap > 0 and current_chunk:
                    overlap_words = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                    current_chunk = overlap_words
                    current_length = sum(len(w) + 1 for w in overlap_words)
                else:
                    current_chunk = []
                    current_length = 0
                
                chunk_index += 1
        
        # Добавляем последний чанк
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append(DocumentChunk(
                text=chunk_text,
                chunk_index=chunk_index,
                metadata={'chunk_size': current_length}
            ))
        
        return chunks
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг PDF файла"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available")
        
        content = {'text': '', 'tables': [], 'pages': [], 'metadata': {}}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Метаданные
                if pdf_reader.metadata:
                    content['metadata'] = {
                        'title': pdf_reader.metadata.get('/Title'),
                        'author': pdf_reader.metadata.get('/Author'),
                        'subject': pdf_reader.metadata.get('/Subject'),
                        'creator': pdf_reader.metadata.get('/Creator'),
                        'producer': pdf_reader.metadata.get('/Producer'),
                        'creation_date': pdf_reader.metadata.get('/CreationDate'),
                        'modification_date': pdf_reader.metadata.get('/ModDate'),
                    }
                    content['title'] = content['metadata'].get('title')
                
                # Текст по страницам
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    content['text'] += page_text + '\n'
                    content['pages'].append({
                        'page_num': page_num + 1,
                        'text': page_text,
                        'rotation': page.get('/Rotate', 0)
                    })
                
        except Exception as e:
            logging.error(f"Error parsing PDF {file_path}: {e}")
            raise
        
        return content
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг DOCX файла"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        content = {'text': '', 'tables': [], 'pages': [], 'metadata': {}}
        
        try:
            doc = Document(file_path)
            
            # Метаданные
            if doc.core_properties:
                content['metadata'] = {
                    'title': doc.core_properties.title,
                    'author': doc.core_properties.author,
                    'subject': doc.core_properties.subject,
                    'created': doc.core_properties.created,
                    'modified': doc.core_properties.modified,
                }
                content['title'] = content['metadata'].get('title')
            
            # Текст
            for paragraph in doc.paragraphs:
                content['text'] += paragraph.text + '\n'
            
            # Таблицы
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append({
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
                
        except Exception as e:
            logging.error(f"Error parsing DOCX {file_path}: {e}")
            raise
        
        return content
    
    def _parse_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг XLSX файла"""
        if not XLSX_AVAILABLE:
            raise ImportError("openpyxl not available")
        
        content = {'text': '', 'tables': [], 'pages': [], 'metadata': {}}
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            # Метаданные
            content['metadata'] = {
                'title': workbook.properties.title,
                'author': workbook.properties.creator,
                'subject': workbook.properties.subject,
                'created': workbook.properties.created,
                'modified': workbook.properties.modified,
            }
            content['title'] = content['metadata'].get('title')
            
            # Данные из всех листов
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content['text'] += f"Sheet: {sheet_name}\n"
                
                # Получаем данные из листа
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data.append([str(cell) if cell is not None else '' for cell in row])
                        content['text'] += ' | '.join([str(cell) if cell is not None else '' for cell in row]) + '\n'
                
                if sheet_data:
                    content['tables'].append({
                        'sheet_name': sheet_name,
                        'data': sheet_data,
                        'rows': len(sheet_data),
                        'cols': len(sheet_data[0]) if sheet_data else 0
                    })
                
        except Exception as e:
            logging.error(f"Error parsing XLSX {file_path}: {e}")
            raise
        
        return content
    
    def _parse_html(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг HTML файла"""
        if not HTML_AVAILABLE:
            raise ImportError("beautifulsoup4 not available")
        
        content = {'text': '', 'tables': [], 'pages': [], 'metadata': {}}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Метаданные
            title_tag = soup.find('title')
            if title_tag:
                content['title'] = title_tag.get_text().strip()
            
            # Убираем скрипты и стили
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Текст
            content['text'] = soup.get_text()
            
            # Таблицы
            for table in soup.find_all('table'):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    content['tables'].append({
                        'data': table_data,
                        'rows': len(table_data),
                        'cols': len(table_data[0]) if table_data else 0
                    })
                
        except Exception as e:
            logging.error(f"Error parsing HTML {file_path}: {e}")
            raise
        
        return content
    
    def _parse_eml(self, file_path: Path) -> Dict[str, Any]:
        """Парсинг EML файла"""
        if not EML_AVAILABLE:
            raise ImportError("email module not available")
        
        content = {'text': '', 'tables': [], 'pages': [], 'metadata': {}}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                eml_content = file.read()
            
            msg = email.message_from_string(eml_content)
            
            # Метаданные
            content['metadata'] = {
                'subject': msg.get('Subject'),
                'from': msg.get('From'),
                'to': msg.get('To'),
                'date': msg.get('Date'),
                'message_id': msg.get('Message-ID'),
            }
            content['title'] = content['metadata'].get('subject')
            
            # Текст
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        content['text'] += part.get_payload(decode=True).decode('utf-8', errors='ignore') + '\n'
            else:
                content['text'] = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                
        except Exception as e:
            logging.error(f"Error parsing EML {file_path}: {e}")
            raise
        
        return content
    
    def _calculate_sha256(self, file_path: Path) -> str:
        """Вычисление SHA256 хеша файла"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256_hash.update(chunk)
        return sha256_hash.hexdigest()
    
    def _get_mime_type(self, extension: str) -> str:
        """Определение MIME типа по расширению"""
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.html': 'text/html',
            '.htm': 'text/html',
            '.eml': 'message/rfc822',
        }
        return mime_types.get(extension, 'application/octet-stream')
    
    def _needs_ocr(self, extension: str, content: Dict[str, Any]) -> bool:
        """Определение необходимости OCR"""
        # PDF может содержать сканы
        if extension == '.pdf':
            text = content.get('text', '')
            # Если текст очень короткий или отсутствует, возможно это сканы
            return len(text.strip()) < 100
        
        # Изображения всегда нуждаются в OCR
        image_extensions = {'.jpg', '.jpeg', '.png', '.tiff', '.bmp'}
        return extension.lower() in image_extensions
