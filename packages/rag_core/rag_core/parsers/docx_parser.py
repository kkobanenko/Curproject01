from __future__ import annotations
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
import hashlib

from .document_parser import DocumentParser

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available, DOCX parsing disabled")

class DOCXParser(DocumentParser):
    """Парсер для DOCX файлов с поддержкой таблиц и метаданных"""
    
    def __init__(self):
        super().__init__()
        self.supported_mime_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Проверяет доступность зависимостей"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available - DOCX parsing will fail")
    
    def can_parse(self, file_path: Path) -> bool:
        """Проверяет, может ли парсер обработать DOCX файл"""
        return (file_path.suffix.lower() == '.docx' and 
                file_path.exists() and 
                DOCX_AVAILABLE)
    
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """Парсит DOCX документ"""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")
        
        # Вычисляем SHA256
        sha256 = self._calculate_sha256(file_path)
        
        # Определяем MIME тип
        mime_type = self.get_mime_type(file_path)
        
        # Парсим содержимое
        content = self._extract_content(file_path)
        
        return {
            'title': content.get('title'),
            'text': content.get('text', ''),
            'tables': content.get('tables', []),
            'pages': content.get('pages', []),
            'sha256': sha256,
            'mime_type': mime_type,
            'needs_ocr': False,  # DOCX не нужен OCR
            'metadata': content.get('metadata', {}),
            'paragraph_count': content.get('paragraph_count', 0),
            'table_count': len(content.get('tables', [])),
            'has_tables': len(content.get('tables', [])) > 0
        }
    
    def _extract_content(self, file_path: Path) -> Dict[str, Any]:
        """Извлекает содержимое DOCX"""
        content = {
            'text': '',
            'tables': [],
            'pages': [],
            'metadata': {},
            'paragraph_count': 0
        }
        
        try:
            doc = Document(file_path)
            
            # Метаданные
            if doc.core_properties:
                content['metadata'] = {
                    'title': doc.core_properties.title,
                    'author': doc.core_properties.author,
                    'subject': doc.core_properties.subject,
                    'created': doc.core_properties.created,
                    'modified': doc.core_properties.modified,
                    'last_modified_by': doc.core_properties.last_modified_by,
                    'revision': doc.core_properties.revision,
                    'category': doc.core_properties.category,
                    'keywords': doc.core_properties.keywords,
                    'comments': doc.core_properties.comments
                }
                content['title'] = content['metadata'].get('title')
            
            # Текст из параграфов
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
                    content['text'] += paragraph.text + '\n'
            
            content['paragraph_count'] = len(paragraphs)
            
            # Таблицы
            for table_index, table in enumerate(doc.tables):
                table_data = self._extract_table_data(table)
                content['tables'].append({
                    'table_index': table_index,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0,
                    'structure': self._analyze_table_structure(table_data)
                })
            
            # Стили и форматирование
            content['styles'] = self._extract_styles(doc)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
        
        return content
    
    def _extract_table_data(self, table) -> List[List[str]]:
        """Извлекает данные из таблицы"""
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Получаем текст из ячейки
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        
        return table_data
    
    def _analyze_table_structure(self, table_data: List[List[str]]) -> Dict[str, Any]:
        """Анализирует структуру таблицы"""
        if not table_data:
            return {}
        
        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0
        
        # Определяем заголовки (первая строка)
        headers = table_data[0] if table_data else []
        
        # Анализируем типы данных в колонках
        column_types = []
        for col_idx in range(cols):
            col_data = [row[col_idx] for row in table_data[1:] if col_idx < len(row)]
            col_type = self._infer_column_type(col_data)
            column_types.append(col_type)
        
        return {
            'headers': headers,
            'column_types': column_types,
            'has_merged_cells': False,  # DOCX не поддерживает объединение ячеек
            'is_empty': all(all(not cell for cell in row) for row in table_data)
        }
    
    def _infer_column_type(self, column_data: List[str]) -> str:
        """Определяет тип данных в колонке"""
        if not column_data:
            return 'empty'
        
        # Проверяем на числа
        numeric_count = 0
        date_count = 0
        
        for value in column_data:
            if value.strip():
                # Проверяем на число
                try:
                    float(value.replace(',', '').replace(' ', ''))
                    numeric_count += 1
                except ValueError:
                    pass
                
                # Простая проверка на дату (можно улучшить)
                if any(separator in value for separator in ['/', '-', '.']):
                    date_count += 1
        
        total_non_empty = len([v for v in column_data if v.strip()])
        
        if total_non_empty == 0:
            return 'empty'
        elif numeric_count / total_non_empty > 0.8:
            return 'numeric'
        elif date_count / total_non_empty > 0.8:
            return 'date'
        else:
            return 'text'
    
    def _extract_styles(self, doc) -> Dict[str, Any]:
        """Извлекает информацию о стилях документа"""
        styles_info = {
            'paragraph_styles': set(),
            'character_styles': set(),
            'table_styles': set()
        }
        
        # Собираем используемые стили
        for paragraph in doc.paragraphs:
            if paragraph.style:
                styles_info['paragraph_styles'].add(paragraph.style.name)
            
            for run in paragraph.runs:
                if run.style:
                    styles_info['character_styles'].add(run.style.name)
        
        for table in doc.tables:
            if table.style:
                styles_info['table_styles'].add(table.style.name)
        
        # Конвертируем в списки
        return {k: list(v) for k, v in styles_info.items()}
    
    def _calculate_sha256(self, file_path: Path) -> str:
        """Вычисляет SHA256 хеш файла"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256_hash.update(chunk)
        return sha256_hash.hexdigest()
